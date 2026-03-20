from datetime import datetime
# coding: utf-8
import psycopg2
import base64 

from os import getcwd
import os
import shutil
from dotenv import load_dotenv
load_dotenv()


from docxtpl import DocxTemplate
from datetime import datetime
# from funciones.cargar_datos_tabla import *
import os
APP_PATH = os.getcwd()

def transformar(dato):
    dato = str(dato)
    # dato = dato.replace("datetime.time(",  " '")
    dato = dato.replace("?",  ' ')
    dato = dato.replace("#",  ' ')
    dato = dato.replace("/",  ' ')
    dato = dato.replace("*",  ' ')
    dato = dato.replace(",",  '-')
    # dato = dato.replace("datetime.datetime(", " '")
    dato = dato.replace("(",  ' ')
    dato = dato.replace("),",  "',")
    dato = dato.replace(")",  ' ')
    dato = dato.replace("$",  ' ')
    dato = dato.replace("%",  ' ')
    dato = dato.replace("{",  ' ')
    dato = dato.replace("}",  ' ')
    dato = dato.replace("[",  ' ')
    dato = dato.replace("]",  ' ')
    dato = dato.replace("<",  ' ')
    dato = dato.replace(">",  ' ')
    dato = dato.replace("¨",  ' ')
    dato = dato.replace("^",  ' ')
    dato = dato.replace("~",  ' ')
    dato = dato.replace("`",  ' ')
    dato = dato.replace("'",  ' ')
    dato = dato.replace('"',  ' ')
    
    
    return dato


def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abri", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    # messsage = "{} de {} del {}".format(day, month, year)

    return [day, month, year]

#Funcion de resultados nueva ayuda
def connect():
    conn = psycopg2.connect(" \
        dbname=dirop \
        user=postgres \
        password=NICval10**")
    return conn


def guardar_evento(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]



    fecha_evento =  datos["fecha_evento"]
    tipo_evento =  datos["tipo_evento"]
    divi_padre =  datos["divi_padre"]
    enemigo =  datos["enemigo"]
    resumen =  datos["resumen"]

    fecha_evento =  fecha_evento.upper()
    tipo_evento =  tipo_evento.upper()
    divi_padre =  divi_padre.upper()
    enemigo =  enemigo.upper()
    resumen =  resumen.upper()

    fecha_evento =  transformar(fecha_evento)
    tipo_evento =  transformar(tipo_evento)
    divi_padre =  transformar(divi_padre)
    enemigo =  transformar(enemigo)
    resumen =  transformar(resumen)
   


    dato="""('{}',	'{}',	'{}',	'{}',	'{}')""".format(fecha_evento, tipo_evento, divi_padre, enemigo, resumen)

    
    query="insert into eventos_relevantes_boletin (fecha_evento, tipo_evento, divi_padre, enemigo, resumen) values"+dato

   
    conn = connect()
    cursor = conn.cursor()
    cursor.execute(query)
    conn.commit()

#1234567890'sdfghjkuytrfvb)(/&%$#) 
                        

    conn.close()
    cursor.close


def planillas(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    tipo_evento = datos["tipo_evento"]
    divi_padre = datos["divi_padre"]
    enemigo = datos["enemigo"]

    fecha_inicio_p = datos["fecha_evento_inicio"]
    fecha_final_p = datos["fecha_evento_final"]

    ti_evento = ""
    ti_division = ""
    ti_enemigo = ""


    if tipo_evento != "":
        ti_evento = " and tipo_evento = '{}'".format(tipo_evento)

    if divi_padre != "":
        ti_division = " and divi_padre = '{}'".format(divi_padre)
        
    if enemigo != "":
        ti_enemigo = " and enemigo = '{}'".format(enemigo)

    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM eventos_relevantes_boletin where fecha_evento >= '{}' and fecha_evento <= '{}' {} {} {} ORDER BY fecha_evento ASC".format(fecha_inicio_p, fecha_final_p, ti_evento, ti_division, ti_enemigo )

    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close

    
    directorio = APP_PATH +"/documentos/plantilla_planillas.docx"


    
    doc = DocxTemplate(directorio)
    # fecha_f = current_date_format(fecha_inicio_p)
    # fecha_i = current_date_format(fecha_final_p)

    fecha = current_date_format(fecha)

    fecha_hoja  = "{} de {} del {}".format(fecha[0], fecha[1], fecha[2])
    # if fecha_i[2] == fecha_f[2]:
    #     mensaje = "del {} de {}  al {} de {} de {}".format(fecha_i[0],fecha_i[1], fecha_f[0],fecha_f[1],fecha_f[2])
    # else:
    #      mensaje = "del {} de {} del {} al {} de {} de {}".format(fecha_i[0],fecha_i[1],fecha_i[2], fecha_f[0],fecha_f[1],fecha_f[2])
    
    # run.text = str(mensaje)

    context = { 
            'fecha_elaboracion':fecha_hoja,
            'fecha_inicio':fecha_inicio_p,
            'fecha_final':fecha_final_p
        }
    
    doc.render(context)
    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 



    LINK = os.getenv('DIRECION_3_B')
    DIRECION = os.getenv('DIRECION_3')

        # hechos_file = contents['hechos_file']
        # leer_resultados = input_json['leer_resultados']

        # print(hechos_file) 
    

    # for files in os.listdir(LINK):
    #     path = os.path.join(LINK, files)
    #     try:
    #         shutil.rmtree(path)
    #     except OSError:
    #         os.remove(path)


    table2 = doc.add_table(rows=1, cols=20) 
    table2.style = 'resultados' 
    row = table2.rows[0].cells 
        

    fila = 1
    row[0].merge(row[0])
    row[0].paragraphs[0].add_run("No.").bold = True

    row[2].merge(row[1])
    row[1].paragraphs[0].add_run("FECHA").bold = True

    row[4].merge(row[3])
    row[3].paragraphs[0].add_run("TIPO EVENTO").bold = True

    row[6].merge(row[5])
    row[5].paragraphs[0].add_run("DIVISION").bold = True
                
    row[8].merge(row[7])
    row[7].paragraphs[0].add_run("ENEMIGO").bold = True

    row[19].merge(row[9])
    row[9].paragraphs[0].add_run("EVENTO").bold = True

    

    for j in data:

        row = table2.add_row().cells 
        row[0].merge(row[0])
        row[0].text = str(fila)

        row[2].merge(row[1])
        row[1].text = str(j[1])

        row[4].merge(row[3])
        row[3].text = str(j[2])

        row[6].merge(row[5])
        row[5].text = str(j[3])
        
        row[8].merge(row[7])
        row[7].text = str(str(j[4])) 

                
        row[19].merge(row[9])
        row[9].text = str(j[5])
        p = row[9].paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
          




        # row[3].merge(row[0])
        # row[0].paragraphs[0].add_run("QUIEN RECIBE BRIGADA").bold = True
   
        # row[9].merge(row[4])
        # row[4].text = str(j[34])+ " - " +str(j[35]+". "+str(j[36]) + " (" + str(j[38])+")" + "  "+ str( j[37]))

        # row[13].merge(row[10])
        # row[10].paragraphs[0].add_run("QUIEN RECIBE BATALLÓN").bold = True
   
        # row[19].merge(row[14])
        # row[14].text = str(j[39])+ " - " +str(j[40]+". "+str(j[41]) + " (" + str(j[43])+")" + "  "+ str( j[42]))
        
        fila =  fila + 1

        parrafo = doc.add_paragraph()


        # row[6].merge(row[1])
        # row[1].text = str(data[j][3])
        # row[7].paragraphs[0].add_run("Tipo").bold = True
        # row[9].merge(row[8])
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()



    doc.save(LINK+ "eventos relevantes.docx")

    
    titulo = "eventos relevantes"

    direcion= DIRECION+str(titulo)+'.docx'
    # print(direcion)
    return [direcion, titulo]

