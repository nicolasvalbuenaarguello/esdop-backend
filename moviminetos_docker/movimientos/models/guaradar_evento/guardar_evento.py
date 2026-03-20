from datetime import datetime
# coding: utf-8
import psycopg2
import base64 

from os import getcwd
import os
import shutil
from dotenv import load_dotenv
load_dotenv()
import docx 
from docx import Document
from docx.shared import Pt
from docxtpl import DocxTemplate
from datetime import datetime
# from funciones.cargar_datos_tabla import *
import os
APP_PATH = os.getcwd()

def transformar(dato):
    dato = str(dato)
    # dato = dato.replace("datetime.time(",  " '")
    dato = dato.replace("?",  ' ')
    dato = dato.replace("#",  ' ')
    dato = dato.replace("/",  ' ')
    dato = dato.replace("*",  ' ')
    dato = dato.replace(",",  ',')
    # dato = dato.replace("datetime.datetime(", " '")
    dato = dato.replace("(",  ' ')
    dato = dato.replace("),",  "',")
    dato = dato.replace(")",  ' ')
    dato = dato.replace("$",  ' ')
    dato = dato.replace("%",  ' ')
    dato = dato.replace("{",  ' ')
    dato = dato.replace("}",  ' ')
    dato = dato.replace("[",  ' ')
    dato = dato.replace("]",  ' ')
    dato = dato.replace("<",  ' ')
    dato = dato.replace(">",  ' ')
    dato = dato.replace("¨",  ' ')
    dato = dato.replace("^",  ' ')
    dato = dato.replace("~",  ' ')
    dato = dato.replace("`",  ' ')
    dato = dato.replace("'",  ' ')
    dato = dato.replace('"',  ' ')
    
    
    return dato
def transformar_cedula(dato):
    dato = str(dato)
    # dato = dato.replace("datetime.time(",  " '")
    dato = dato.replace("?",  ' ')
    dato = dato.replace("#",  ' ')
    dato = dato.replace("/",  ' ')
    dato = dato.replace("*",  ' ')
    dato = dato.replace(",",  ',')
    dato = dato.replace(".",  '')
    # dato = dato.replace("datetime.datetime(", " '")
    dato = dato.replace("(",  ' ')
    dato = dato.replace("),",  "',")
    dato = dato.replace(")",  ' ')
    dato = dato.replace("$",  ' ')
    dato = dato.replace("%",  ' ')
    dato = dato.replace("{",  ' ')
    dato = dato.replace("}",  ' ')
    dato = dato.replace("[",  ' ')
    dato = dato.replace("]",  ' ')
    dato = dato.replace("<",  ' ')
    dato = dato.replace(">",  ' ')
    dato = dato.replace("¨",  ' ')
    dato = dato.replace("^",  ' ')
    dato = dato.replace("~",  ' ')
    dato = dato.replace("`",  ' ')
    dato = dato.replace("'",  ' ')
    dato = dato.replace('"',  ' ')
    
    
    return dato
#Funcion de resultados nueva ayuda
def connect():
    conn = psycopg2.connect(" \
        dbname=dirop \
        user=postgres \
        password=NICval10**")
    return conn



def eliminar_evento(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    id_coordinador = datos["id_coordinador"]




    # print(documento)

    # image_64_encode = base64.encode(documento)
    query="""DELETE FROM coordinadores WHERE id = {}""".format(id_coordinador)

    conn = connect()
    cursor = conn.cursor()
    cursor.execute(query)
    conn.commit()

                
    query_eventos =  "SELECT * FROM coordinadores ORDER BY id ASC"

            
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close



    return [data]

def guardar_evento(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    unidad_solitante = datos["unidad_solitante"]

    unidad_solitante =  unidad_solitante.upper()



    # print(documento)

    # image_64_encode = base64.encode(documento)

    dato="""('{}')""".format(unidad_solitante)

    
    query="insert into UNIDAD_SOLICITANTE (unidad) values"+dato

    conn = connect()
    cursor = conn.cursor()
    cursor.execute(query)
    conn.commit()

                
    query_eventos =  "SELECT * FROM UNIDAD_SOLICITANTE ORDER BY unidad ASC"

            
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close
    return [data]


def guardar_movimiento_unidad(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    hr_movimiento = datos["hr_movimiento"]
    fecha_inicio = datos["fecha_inicio"]
    fecha_termino = datos["fecha_termino"]
    unidad_solicitante = datos["unidad_solicitante"]
    tipo_escolta = datos["tipo_escolta"]
    ruta_completa = datos["ruta_completa"]
    motivos_escolta = datos["motivos_escolta"]
    unidades_comprometidas = datos["unidades_comprometidas"]
    hr_unidad = datos["hr_unidad"]
    cantidad_vehiculos = datos["cantidad_vehiculos"]
    cantidad_faces = datos["cantidad_faces"]
    encargado_material = datos["encargado_material"]
    telefono = datos["telefono"]

    hr_movimiento =  transformar(hr_movimiento)
    fecha_inicio =  transformar(fecha_inicio)
    fecha_termino =  transformar(fecha_termino)
    unidad_solicitante =  transformar(unidad_solicitante)
    tipo_escolta =  transformar(tipo_escolta)
    ruta_completa =  transformar(ruta_completa)
    motivos_escolta =  transformar(motivos_escolta)
    unidades_comprometidas =  transformar(unidades_comprometidas)
    hr_unidad =  transformar(hr_unidad)
    cantidad_vehiculos =  transformar(cantidad_vehiculos)
    cantidad_faces =  transformar(cantidad_faces)
    encargado_material =  transformar(encargado_material)
    telefono =  transformar(telefono)


    hr_movimiento = hr_movimiento.upper()
    fecha_inicio = fecha_inicio.upper()
    fecha_termino = fecha_termino.upper()
    unidad_solicitante = unidad_solicitante.upper()
    tipo_escolta = tipo_escolta.upper()
    ruta_completa = ruta_completa.upper()
    motivos_escolta = motivos_escolta.upper()
    unidades_comprometidas = unidades_comprometidas.upper()
    hr_unidad = hr_unidad.upper()
    cantidad_vehiculos = cantidad_vehiculos.upper()
    cantidad_faces = cantidad_faces.upper()
    encargado_material = encargado_material.upper()
    telefono = telefono.upper()

    estado_escolta = " DESARROLLO"


    # print(documento)

    # image_64_encode = base64.encode(documento)

    

    dato="""('{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}')""".format(hr_movimiento, fecha_inicio, fecha_termino, unidad_solicitante, tipo_escolta, ruta_completa, motivos_escolta, unidades_comprometidas, hr_unidad, cantidad_vehiculos, cantidad_faces, encargado_material, telefono, estado_escolta)

    
    query="insert into PLANILLA_MOVIMIENTO (hr_movimiento, fecha_inicio, fecha_termino, unidad_solicitante, tipo_escolta, ruta_completa, motivos_escolta, unidades_comprometidas, hr_unidad, cantidad_vehiculos, cantidad_faces, encargado_material, telefono, estado_escolta) values"+dato

    conn = connect()
    cursor = conn.cursor()
    cursor.execute(query)
    conn.commit()
    conn.close()
    cursor.close




def listado_movientos(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    fecha_inicio = datos["fecha_inicio"]
    fecha_termino = datos["fecha_termino"]

    fecha_1 = ""
    if fecha_inicio !="" and fecha_termino !="":
        fecha_1 = "where fecha_inicio >= '{}' and fecha_termino <= '{}'".format(fecha_inicio, fecha_termino)
    elif fecha_inicio !="" and fecha_termino =="":
        fecha_1 = "where fecha_inicio >= '{}' ".format(fecha_inicio)
    elif fecha_inicio =="" and fecha_termino !="":
        fecha_1 = "where fecha_termino <= '{}' ".format(fecha_termino)



    conn = connect()
    cursor = conn.cursor()

                
    query_eventos =  "SELECT * FROM PLANILLA_MOVIMIENTO {} ORDER BY fecha_inicio ASC".format(fecha_1)

            
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close
    return [data]


def listado_unidades_solictantes(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    # print(documento)

    # image_64_encode = base64.encode(documento)



    conn = connect()
    cursor = conn.cursor()


                
    query_eventos =  "SELECT * FROM UNIDAD_SOLICITANTE ORDER BY unidad ASC"

            
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close
    return [data]


def editar_evento(datos):

    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    id_coordinador = datos["id_coordinador"]

    grd_coordinador = datos["grd_coordinador"]
    nombre_coordinador = datos["nombre_coordinador"]
    cargo_coordinador = datos["cargo_coordinador"]

    sigla_unidad = datos["sigla_unidad"]
    nombre_unidad = datos["nombre_unidad"]
    tel_coordinador= datos["tel_coordinador"]

    grd_coordinador =  grd_coordinador.upper()
    nombre_coordinador =  nombre_coordinador.upper()
    cargo_coordinador =  cargo_coordinador.upper()
    sigla_unidad =  sigla_unidad.upper()
    nombre_unidad =  nombre_unidad.upper()
    tel_coordinador =  tel_coordinador.upper()


    # print(documento)

    # image_64_encode = base64.encode(documento)

    query="""UPDATE coordinadores SET grd_coordinador='{}', nombre_coordinador='{}',cargo_coordinador='{}',sigla_unidad='{}',nombre_unidad='{}',tel_coordinador='{}' WHERE id={}""".format(grd_coordinador, nombre_coordinador, cargo_coordinador, sigla_unidad, nombre_unidad, tel_coordinador, id_coordinador)

    # dato="""('{}',	'{}',	'{}',	'{}',	'{}',	'{}')""".format(grd_coordinador,	nombre_coordinador,	cargo_coordinador,	sigla_unidad,	nombre_unidad,	tel_coordinador)

    
    # query="insert into coordinadores (grd_coordinador, nombre_coordinador, cargo_coordinador,	sigla_unidad,	nombre_unidad,	tel_coordinador ) values"+dato

    conn = connect()
    cursor = conn.cursor()
    cursor.execute(query)
    conn.commit()

                
    query_eventos =  "SELECT * FROM coordinadores ORDER BY id ASC"

            
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close



    return [data]


from docx.shared import Pt
from docxtpl import DocxTemplate
import os
import psycopg2

def crear_planilla(datos):
    """
    Crea un documento Word basado en la plantilla y los datos de PLANILLA_MOVIMIENTO.
    """
    # Datos de entrada
    permiso = datos.get("permiso")
    unidad = datos.get("unidad")
    numero_planilla = datos.get("numero_planilla")
    fecha_planilla = datos.get("fecha_planilla")
    listados_orfeo = datos.get("listados_orfeo", "")

    # Convertir los códigos en lista y limpiar espacios
    orfeos = [x.strip() for x in listados_orfeo.split(",") if x.strip()]
    if not orfeos:
        raise ValueError("No se enviaron códigos de hr_movimiento válidos")

    # Preparar placeholders para la consulta
    placeholders = ','.join(['%s'] * len(orfeos))
    query_eventos = f"""
        SELECT * 
        FROM PLANILLA_MOVIMIENTO 
        WHERE hr_movimiento IN ({placeholders}) 
        ORDER BY fecha_inicio ASC
    """

    # Conexión a la base de datos
    conn = connect()  # Tu función connect() ya existente
    cursor = conn.cursor()
    cursor.execute(query_eventos, orfeos)
    data = cursor.fetchall()
    cursor.close()
    conn.close()

    # Cargar plantilla Word
    directorio = APP_PATH + "/documentos/template.docx"
    doc_p = DocxTemplate(directorio)
    
    context = { 
        'numero': numero_planilla,
        'fecha': fecha_planilla
    }
    doc_p.render(context)

    # Configurar fuente
    style = doc_p.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Crear tabla
    table2 = doc_p.add_table(rows=1, cols=20)
    table2.style = 'Table Grid'  # Estilo válido de tabla
    row = table2.rows[0].cells

    # Encabezados
    row[0].merge(row[0])
    row[0].paragraphs[0].add_run("No.").bold = False

    row[2].merge(row[1])
    row[1].paragraphs[0].add_run("GDO").bold = False

    row[11].merge(row[3])
    row[3].paragraphs[0].add_run("APELLIDOS Y NOMBRE").bold = False

    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("CC").bold = False

    row[19].merge(row[16])
    row[16].paragraphs[0].add_run("UNIDAD").bold = False

    # Rellenar datos
    for fila, j in enumerate(data, start=1):
        row = table2.add_row().cells
        row[0].merge(row[0])
        row[0].text = str(fila)

        row[2].merge(row[1])
        row[1].text = str(j[10])

        row[11].merge(row[3])
        row[3].text = str(j[13])

        row[15].merge(row[12])
        row[12].text = str(j[12])

        row[19].merge(row[16])
        datos_unidad = str(j[7]).replace("N A", '').replace("N/A", '')
        row[16].text = datos_unidad

    # Guardar documento
    documento = f"{numero_planilla}_planilla"
    LINK = os.getenv('DIRECION_3_B')
    DIRECION = os.getenv('DIRECION_3')
    doc_p.save(LINK + documento + ".docx")

    direcion = DIRECION + str(documento) + '.docx'
    return [direcion, documento]
