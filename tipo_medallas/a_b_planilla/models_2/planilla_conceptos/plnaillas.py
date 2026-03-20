from datetime import datetime
import docx 
from docx import Document
from docx.shared import Pt
# coding: utf-8
import psycopg2
import base64 

from os import getcwd
import os
import shutil
from dotenv import load_dotenv
load_dotenv()

from docxtpl import DocxTemplate
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, colors, fills, Alignment, PatternFill, NamedStyle
# from funciones.cargar_datos_tabla import *
import os
APP_PATH = os.getcwd()

#formato de fechas 
def valorgrado(grd):
    valor = 0
    if grd == "GR":
        valor = 0
    elif grd == "MG":
        valor = 1
    elif grd == "BG":
        valor = 2
    elif grd == "CR":
        valor = 3
    elif grd == "TC":
        valor = 4
    elif grd == "MY":
        valor = 5
    elif grd == "CT":
        valor = 6
    elif grd == "TE":
        valor = 7
    elif grd == "ST":
        valor = 8
    elif grd == "SMC":
        valor = 9
    elif grd == "SM":
        valor = 10
    elif grd == "SP":
        valor = 11
    elif grd == "SV":
        valor = 12
    elif grd == "SS":
        valor = 13
    elif grd == "CP":
        valor = 14
    elif grd == "CS":
        valor = 15
    elif grd == "C3":
        valor = 16
    elif grd == "DGSLP":
        valor = 17
    elif grd == "SLP":
        valor = 18
    elif grd == "DG":
        valor = 19
    elif grd == "SL18":
        valor = 20
    elif grd == "SL12":
        valor = 21
    elif grd == "----":
        valor = 22
    return valor

def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abri", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    # messsage = "{} de {} del {}".format(day, month, year)

    return [day, month, year]
#Funcion de resultados nueva ayuda
def connect():
    conn = psycopg2.connect(" \
        dbname=dirop \
        user=postgres \
        password=NICval10**")
    return conn


def informe_entrega(datos):
    #try:
            
        # fecha = fecha.strftime('%d-%m-%Y ')
        #datos
        
        
        permiso = datos["permiso"]
        unidad = datos["unidad"]
        proyecto = datos["proyecto"]
        medalla = datos["medalla"]

        conn = connect()
        cursor = conn.cursor()

        query_eventos =  "SELECT * FROM medallas where proyecto = '{}' and medalla = '{}' ".format(proyecto, medalla)
       
        cursor.execute(query_eventos)
        data = cursor.fetchall()
        conn.close()
        cursor.close

        directorio = APP_PATH +"/a_b_planilla/documentos/templae_listado.docx"

        fecha = datetime.today()
        doc_p = DocxTemplate(directorio)
            # fecha_f = current_date_format(fecha_inicio_p)
            # fecha_i = current_date_format(fecha_final_p)

        fecha = current_date_format(fecha)

        fecha_hoja  = "{} de {} del {}".format(fecha[0], fecha[1], fecha[2])
            # if fecha_i[2] == fecha_f[2]:
            #     mensaje = "del {} de {}  al {} de {} de {}".format(fecha_i[0],fecha_i[1], fecha_f[0],fecha_f[1],fecha_f[2])
            # else:
            #      mensaje = "del {} de {} del {} al {} de {} de {}".format(fecha_i[0],fecha_i[1],fecha_i[2], fecha_f[0],fecha_f[1],fecha_f[2])

            # run.text = str(mensaje)

        tipo_medalla = ""
        tipo_medalla_upper=""
        if medalla == "MEDALLA DESEMPEÑO OPERACIONAL":
                msm ="""Dentro de mencionada operación militar,  se respeto la vida y se garantizo la seguridad, fundamentado en el respeto de los Derechos Humanos y el acatamiento al Derecho Internacional Humanitario."""

                tipo_medalla = "“Servicios distinguidos en Orden Público”"
                tipo_medalla_upper = tipo_medalla.upper()
        elif medalla == "MEDALLA AL VALOR":
   
                tipo_medalla = "“Al Valor”."
                tipo_medalla_upper = tipo_medalla.upper()
        elif medalla == "MEDALLA HERIDOS EN ACCIÓN":

                tipo_medalla = "“Herido en Acción”"
                tipo_medalla_upper = tipo_medalla.upper()
        elif medalla == "MEDALLA A LA BANDERA DE GUERRA":

                tipo_medalla = "“MEDALLA A LA BANDERA DE GUERRA”"
                tipo_medalla_upper = tipo_medalla.upper()
        else:
     
                tipo_medalla = ""


        context = { 
                    'fecha_elaboracion':fecha_hoja,
                    'tipo_medalla':tipo_medalla,
                    'tipo_medalla_upper':tipo_medalla_upper
                }

        doc_p.render(context)

        style = doc_p.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        table2 = doc_p.add_table(rows=1, cols=20) 
        table2.style = 'listado_personal_medallas' 
        row = table2.rows[0].cells 
            
        fila = 1
        row[0].merge(row[0])
        row[0].paragraphs[0].add_run("No.").bold = False
   

        row[2].merge(row[1])
        row[1].paragraphs[0].add_run("GDO").bold = False

        row[11].merge(row[3])
        row[3].paragraphs[0].add_run("APELLIDOS Y NOMBRE").bold = False

        row[15].merge(row[12])
        row[12].paragraphs[0].add_run("CC").bold = False
                    
        row[19].merge(row[16])
        row[16].paragraphs[0].add_run("UNIDAD").bold = False

        datos = []
        dato_s=[]
        
        for x in data:
            grd  =  valorgrado(x[10])
   
            dato_s =list(x)
            dato_s.append(grd)
            datos.append(tuple(dato_s))
            
        data = sorted(datos, key=lambda d: d[-1], reverse=False)

        for j in data:

            row = table2.add_row().cells 
            row[0].merge(row[0])
            row[0].text = str(fila)

            row[2].merge(row[1])
            row[1].text = str(j[10])

            row[11].merge(row[3])
            row[3].text = str(j[13])

            row[15].merge(row[12])
            row[12].text = str(j[12])

            row[19].merge(row[16])
            datos = str(j[7]).replace("N A", '')
            datos = datos.replace("N/A", '')
            row[16].text = str(datos)
            fila = fila +1

        #doc_p.add_heading('GeeksForGeeks', 0) 




                # hechos_file = contents['hechos_file']
                # leer_resultados = input_json['leer_resultados']

                # print(hechos_file) 


            # for files in os.listdir(LINK):
            #     path = os.path.join(LINK, files)
            #     try:
            #         shutil.rmtree(path)
            #     except OSError:
            #         os.remove(path)

            
        LINK = os.getenv('DIRECION_3_B')
        DIRECION = os.getenv('DIRECION_3')
        doc_p.save(LINK+ "listado_oficio.docx")

        
        titulo = "listado_oficio"

        direcion= DIRECION+str(titulo)+'.docx'
        # print(direcion)
        return [direcion, titulo]
    #except :
        return["error", "error"]

def planillas(datos):

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]
    proyecto = datos["proyecto"]
    medalla = datos["medalla"]
    

    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM medallas where proyecto = '{}' and medalla =  '{}' order by id asc ".format(proyecto, medalla)
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close

    directorio = APP_PATH +"/a_b_planilla/documentos/temple_orden_publico.docx"

    for x in data:
        nombre_doc = x[0]
        fecha = datetime.today()
        doc = DocxTemplate(directorio)
        # fecha_f = current_date_format(fecha_inicio_p)
        # fecha_i = current_date_format(fecha_final_p)

        fecha = current_date_format(fecha)

        fecha_hoja  = "{} de {} del {}".format(fecha[0], fecha[1], fecha[2])
        # if fecha_i[2] == fecha_f[2]:
        #     mensaje = "del {} de {}  al {} de {} de {}".format(fecha_i[0],fecha_i[1], fecha_f[0],fecha_f[1],fecha_f[2])
        # else:
        #      mensaje = "del {} de {} del {} al {} de {} de {}".format(fecha_i[0],fecha_i[1],fecha_i[2], fecha_f[0],fecha_f[1],fecha_f[2])

        # run.text = str(mensaje)
        msm =""
        resumen=""
        no_resultados = ""
        tipo_medalla = ""
        como = ""
        if x[2] == "MEDALLA DESEMPEÑO OPERACIONAL":
            msm ="""Dentro de mencionada operación militar,  se respeto la vida y se garantizo la seguridad, fundamentado en el respeto de los Derechos Humanos y el acatamiento al Derecho Internacional Humanitario."""
            resumen = x[23]
            tipo_medalla = "“Servicios distinguidos en Orden Público”."
            como =  x[16]
        elif x[2] == "MEDALLA AL VALOR":
            no_resultados = x[23]
            tipo_medalla = "“Al Valor”."
            como =  "“Al Valor”."
        elif x[2] == "MEDALLA HERIDOS EN ACCIÓN":
            no_resultados = x[23]
            tipo_medalla = "“Herido en Acción”."
            if x[30]:
                como = x[30]

        elif x[2] == "MEDALLA A LA BANDERA DE GUERRA":
            no_resultados = x[23]
            tipo_medalla = "“Servicios distinguidos en Orden Público”."
            como = "“Servicios distinguidos en Orden Público”."
        else:
            no_resultados = ""
            tipo_medalla = ""
            como=""

        unidad = x[8] +"-"+ x[9] 

        if x[20] != '':
            lugar = "Municipio "+x[20] + " ("+ x[19]+")"
        else: 
             lugar = x[31] + " ("+ x[19]+")"

        fecha_hechos = x[17]
        d = datetime.strptime(fecha_hechos, '%Y-%m-%d')
        d = d.strftime('%d/%m/%Y') # Para dar formato

        fecha =""
        fecha_hechos_final = x[18]
        try:
            if fecha_hechos_final != '':
                d_2 = datetime.strptime(fecha_hechos_final, '%Y-%m-%d')
                d_2 = d_2.strftime('%d/%m/%Y') # Para dar formato

            
                if fecha_hechos == fecha_hechos_final:
                    fecha  = d
                else:
                    fecha = d + " - " +d_2
            else:
                fecha = d 
        except:
            fecha = d 

        context = { 
                'fecha_elaboracion':fecha_hoja,
                'radicado':x[6],
                'grd':x[10],
                'arm':x[11],
                'nombre_apellido':x[13],
                'cc':x[12],
                'unidad':unidad,
                'operacion':"Operación "+x[15],
                'fecha_hecho':fecha,
                'dep_mun_ver':lugar,
                'tipo_solicitud':como,
                'amenaza':x[22],
                'resultados':resumen,
                'no_resultados':no_resultados,
                'justicificacion_orden_publico':msm,
                'tipo_medalla':tipo_medalla
            }

        doc.render(context)
        import docx 

        doc2 = docx.Document() 

        doc2.add_heading('GeeksForGeeks', 0) 




            # hechos_file = contents['hechos_file']
            # leer_resultados = input_json['leer_resultados']

            # print(hechos_file) 


        # for files in os.listdir(LINK):
        #     path = os.path.join(LINK, files)
        #     try:
        #         shutil.rmtree(path)
        #     except OSError:
        #         os.remove(path)

        directorio_guardar = APP_PATH +"/a_b_planilla/documentos_conceptos/"+str(nombre_doc)+".docx"

        doc.save(directorio_guardar)

    #filename_master is name of the file you want to merge the docx file into

def combine_all_docx(datos):
    try:
        medalla = datos["medalla"]
        dirercion_archvios = APP_PATH +"/a_b_planilla/documentos_conceptos/"
        listaPdfs = os.listdir(dirercion_archvios)

        from docxcompose.composer import Composer
        from docx import Document as Document_compose
        number_of_sections=len(listaPdfs)

        master = Document_compose(dirercion_archvios+listaPdfs[0])
        composer = Composer(master)
        for i in range(1, number_of_sections):

            doc_temp = Document_compose(dirercion_archvios+listaPdfs[i])
            composer.append(doc_temp)

        #composer.save("combined_file.docx")
        
        LINK = os.getenv('DIRECION_3_B')
        DIRECION = os.getenv('DIRECION_3')
        composer.save(LINK+"conceptos_"+medalla+"_.docx")

        
        titulo = "conceptos_"+medalla+"_"

        direcion= DIRECION+str(titulo)+'.docx'
        # print(direcion)
        return [direcion, titulo]
    except :
        return["error", "error"]


def validacion_cantidad_tupla(tupla):
                  
    if len(tupla) == 1:
        valor = " = '"+tupla[0]+"'"
        
    else:
        tupla = tuple(tupla)
        valor = " in {}" .format(tupla)
    return valor 
def excel(datos):

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    filtro_division = datos["filtro_division"]
    filtro_medallas = datos["filtro_medallas"]
    filtro_tipo = datos["filtro_tipo"]

    
    filtro_division=filtro_division.split(",")
    filtro_medallas=filtro_medallas.split(",")
    filtro_tipo=filtro_tipo.split(",")

    valor_2 = validacion_cantidad_tupla(filtro_division)
    valor_3 = validacion_cantidad_tupla(filtro_medallas)
    valor_4 = validacion_cantidad_tupla(filtro_tipo)

    division =""
    medalla = ""
    tipo = ""
    
    if len(filtro_division[0])>1:
          division = "where division {}".format(valor_2)


    if len(filtro_medallas[0])>1:
          medalla = " and medalla {}".format(valor_3)

     
    if len(filtro_tipo[0])>1:
          tipo = "and estado_medalla {}".format(valor_4)

    if division =="":
        if medalla !="":
            medalla =   medalla.replace("and", "where")
        
        elif tipo !="":
            tipo =  tipo.replace("and", "where")

         



    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM medallas {} {} {} order by id asc ".format(division, medalla, tipo)
    print(query_eventos)
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close

    
    wb = openpyxl.Workbook()


                
    BUILTIN_FORMATS = {
                                0: 'General',
                                1: '0',
                                2: '0.00',
                                3: '#,##0',
                                4: '#,##0.00',
                                5: '"$"#,##0_);("$"#,##0)',
                                6: '"$"#,##0_);[Red]("$"#,##0)',
                                7: '"$"#,##0.00_);("$"#,##0.00)',
                                8: '"$"#,##0.00_);[Red]("$"#,##0.00)',
                                9: '0%',
                                10: '0.00%',
                                11: '0.00E+00',
                                12: '# ?/?',
                                13: '# ??/??',
                                14: 'mm-dd-yy',
                                15: 'd-mmm-yy',
                                16: 'd-mmm',
                                17: 'mmm-yy',
                                18: 'h:mm AM/PM',
                                19: 'h:mm:ss AM/PM',
                                20: 'h:mm',
                                21: 'h:mm:ss',
                                22: 'm/d/yy h:mm',

                                37: '#,##0_);(#,##0)',
                                38: '#,##0_);[Red](#,##0)',
                                39: '#,##0.00_);(#,##0.00)',
                                40: '#,##0.00_);[Red](#,##0.00)',

                                41: r'_(* #,##0_);_(* \(#,##0\);_(* "-"_);_(@_)',
                                42: r'_("$"* #,##0_);_("$"* \(#,##0\);_("$"* "-"_);_(@_)',
                                43: r'_(* #,##0.00_);_(* \(#,##0.00\);_(* "-"??_);_(@_)',

                                44: r'_("$"* #,##0.00_)_("$"* \(#,##0.00\)_("$"* "-"??_)_(@_)',
                                45: 'mm:ss',
                                46: '[h]:mm:ss',
                                47: 'mmss.0',
                                48: '##0.0E+0',
                                49: '@', }
            
    frmt = BUILTIN_FORMATS[3] 
    frmt_d = BUILTIN_FORMATS[4] 


    nombre_hoja = "base de datos"

    hoja = wb.active       
    hoja.title = nombre_hoja
    numero  =  1

    hoja.cell(row=numero, column=1, value=str("No"))
    hoja.cell(row=numero, column=2, value=str("Recibida por:"))
    hoja.cell(row=numero, column=3, value=str("Medalla"))
    hoja.cell(row=numero, column=4, value=str("Fecha que recibió"))
    hoja.cell(row=numero, column=5, value=str("Mes"))
    hoja.cell(row=numero, column=6, value=str("Oficio"))
    hoja.cell(row=numero, column=7, value=str('Radicado'))

    hoja.cell(row=numero, column=8, value=str("Batallón"))
    hoja.cell(row=numero, column=9, value=str("Brigada"))
    hoja.cell(row=numero, column=10, value=str("Division"))

    hoja.cell(row=numero, column=11, value=str("Grd"))
    hoja.cell(row=numero, column=12, value=str("Arma"))
    hoja.cell(row=numero, column=13, value=str("Cedula"))
    hoja.cell(row=numero, column=14, value=str('Nombres Apellidos'))

    hoja.cell(row=numero, column=15, value=str('Categoria'))
    hoja.cell(row=numero, column=16, value=str('Operacion'))
    hoja.cell(row=numero, column=17, value=str('Motivo'))

    hoja.cell(row=numero, column=18, value=str('Fecha Hecho'))
    hoja.cell(row=numero, column=19, value=str('Fecha Hecho Final'))
    hoja.cell(row=numero, column=20, value=str('Departamento'))
    hoja.cell(row=numero, column=21, value=str('Municipio'))
    hoja.cell(row=numero, column=22, value=str('Veredad'))

    hoja.cell(row=numero, column=23, value=str('Amenaza'))
    hoja.cell(row=numero, column=24, value=str('Resumen'))
    hoja.cell(row=numero, column=25, value=str('Entrega Medalla DIFAB - DINEG'))
    hoja.cell(row=numero, column=26, value=str('Estado Medalla'))

    hoja.cell(row=numero, column=27, value=str('Resolucion'))
    hoja.cell(row=numero, column=28, value=str('Recreto'))
    hoja.cell(row=numero, column=29, value=str('Recha Resolucion'))
    hoja.cell(row=numero, column=30, value=str('Proyecto'))
    hoja.cell(row=numero, column=31, value=str('Como'))

    hoja.cell(row=numero, column=32, value=str('Area de Responsabilidad'))
    hoja.cell(row=numero, column=33, value=str('Observaciones'))

    numero = numero +1


    num = 0
    for x in data:
        num = num + 1  
        hoja.cell(row=numero, column=1, value=str(num))
        hoja.cell(row=numero, column=2, value=str(x[1]))
        hoja.cell(row=numero, column=3, value=str(x[2]))
        hoja.cell(row=numero, column=4, value=str(x[3]))
        hoja.cell(row=numero, column=5, value=str(x[4]))
        hoja.cell(row=numero, column=6, value=str(x[5]))
        hoja.cell(row=numero, column=7, value=str(x[6]))
        hoja.cell(row=numero, column=8, value=str(x[7]))
        hoja.cell(row=numero, column=9, value=str(x[8]))
        hoja.cell(row=numero, column=10, value=str(x[9]))
        hoja.cell(row=numero, column=11, value=str(x[10]))
        hoja.cell(row=numero, column=12, value=str(x[11]))
        hoja.cell(row=numero, column=13, value=str(x[12]))
        hoja.cell(row=numero, column=14, value=str(x[13]))
        hoja.cell(row=numero, column=15, value=str(x[14]))
        hoja.cell(row=numero, column=16, value=str(x[15]))
        hoja.cell(row=numero, column=17, value=str(x[16]))
        hoja.cell(row=numero, column=18, value=str(x[17]))
        hoja.cell(row=numero, column=19, value=str(x[18]))

        hoja.cell(row=numero, column=20, value=str(x[19]))
        hoja.cell(row=numero, column=21, value=str(x[20]))
        hoja.cell(row=numero, column=22, value=str(x[21]))
        hoja.cell(row=numero, column=23, value=str(x[22]))
        hoja.cell(row=numero, column=24, value=str(x[23]))

        hoja.cell(row=numero, column=25, value=str(x[24]))
        hoja.cell(row=numero, column=26, value=str(x[25]))
        hoja.cell(row=numero, column=27, value=str(x[26]))
        hoja.cell(row=numero, column=28, value=str(x[27]))
        hoja.cell(row=numero, column=29, value=str(x[28]))
        hoja.cell(row=numero, column=30, value=str(x[29]))
        hoja.cell(row=numero, column=31, value=str(x[30]))

        hoja.cell(row=numero, column=32, value=str(x[31]))
        hoja.cell(row=numero, column=33, value=str(x[32]))
        
        numero = numero + 1






    for cell in hoja["1:1"]: 
        cell.font = Font(bold=True, size=12, color='FFFFFF')
        cell.fill  = PatternFill('solid', start_color="800000")
            # cell.alignment = Alignment(wrap_text=True)
        cell.alignment=Alignment(
            horizontal='center',
            vertical='top',
            text_rotation=0,
            wrap_text=False,
            shrink_to_fit=False,
            indent=0)




        
    LINK = os.getenv('DIRECION_3_B')
    DIRECION = os.getenv('DIRECION_3')

    wb.save(LINK+ str("listado medallas")+'.xlsx')
    wb.close()

    # titulo = "resultados UT"

    direcion= DIRECION+str("listado medallas")+'.xlsx'

    return [direcion, "listado medallas"]

    #filename_master is name of the file you want to merge the docx file into
