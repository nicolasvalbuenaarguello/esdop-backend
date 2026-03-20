from datetime import datetime
# coding: utf-8
import psycopg2
import base64 

from os import getcwd
import os
import shutil
from dotenv import load_dotenv
load_dotenv()

from docxtpl import DocxTemplate
from datetime import datetime
# from funciones.cargar_datos_tabla import *
import os
APP_PATH = os.getcwd()

#formato de fechas 

def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abri", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    # messsage = "{} de {} del {}".format(day, month, year)

    return [day, month, year]
#Funcion de resultados nueva ayuda
def connect():
    conn = psycopg2.connect(" \
        dbname=dirop \
        user=postgres \
        password=NICval10**")
    return conn

def info_alertas_filtro(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos

    tipo_filtro = datos["tipo_filtro"]
    filtro = datos["filtro"]

    
    


    if tipo_filtro != "" and filtro !="":
        print(tipo_filtro)
        print(filtro)
        query_eventos =  "SELECT * FROM alertas where {} like '%{}%' ORDER BY FECHA ASC".format(tipo_filtro, filtro)
    else:
        query_eventos =  "SELECT * FROM alertas ORDER BY FECHA ASC"

    conn = connect()
    cursor = conn.cursor()

    
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close


 
    # print(direcion)
    return [data, filtro]


def info_mapa(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    fecha_inicio_p = datos["fecha_inicio_m"]
    fecha_final_p = datos["fecha_final_m"]

    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM alertas where fecha >= '{}' and fecha <= '{}' ORDER BY FECHA ASC".format(fecha_inicio_p, fecha_final_p)
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close


 
    # print(direcion)
    return [data]


def info_mapa_division(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    sigla_unidad = datos["sigla_unidad"]

    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM alertas where destino = '{}' ORDER BY FECHA ASC".format(sigla_unidad)
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close


 
    # print(direcion)
    return [data]
def info_mapa_fecha(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    fecha = datos["fecha"]

    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM alertas where fecha = '{}' ORDER BY FECHA ASC".format(fecha)
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close


 
    # print(direcion)
    return [data]

def planillas(datos):
    
    fecha = datetime.today()

    # fecha = fecha.strftime('%d-%m-%Y ')
    #datos
    permiso = datos["permiso"]
    unidad = datos["unidad"]

    fecha_inicio_p = datos["fecha_inicio_p"]
    fecha_final_p = datos["fecha_final_p"]
    
    oficial_coe  = datos["oficial_coe"]
    oficial_coe = oficial_coe.upper()
    conn = connect()
    cursor = conn.cursor()

    query_eventos =  "SELECT * FROM alertas where fecha >= '{}' and fecha <= '{}' ORDER BY FECHA ASC".format(fecha_inicio_p, fecha_final_p)
     
    cursor.execute(query_eventos)
    data = cursor.fetchall()
    conn.close()
    cursor.close

    
    directorio = APP_PATH +"/documentos/plantilla_planillas.docx"


    
    doc = DocxTemplate(directorio)
    # fecha_f = current_date_format(fecha_inicio_p)
    # fecha_i = current_date_format(fecha_final_p)

    fecha = current_date_format(fecha)

    fecha_hoja  = "{} de {} del {}".format(fecha[0], fecha[1], fecha[2])
    # if fecha_i[2] == fecha_f[2]:
    #     mensaje = "del {} de {}  al {} de {} de {}".format(fecha_i[0],fecha_i[1], fecha_f[0],fecha_f[1],fecha_f[2])
    # else:
    #      mensaje = "del {} de {} del {} al {} de {} de {}".format(fecha_i[0],fecha_i[1],fecha_i[2], fecha_f[0],fecha_f[1],fecha_f[2])
    
    # run.text = str(mensaje)

    context = { 
            'fecha_elaboracion':fecha_hoja
        }
    
    doc.render(context)
    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 



    LINK = os.getenv('DIRECION_3_B')
    DIRECION = os.getenv('DIRECION_3')

        # hechos_file = contents['hechos_file']
        # leer_resultados = input_json['leer_resultados']

        # print(hechos_file) 
    

    # for files in os.listdir(LINK):
    #     path = os.path.join(LINK, files)
    #     try:
    #         shutil.rmtree(path)
    #     except OSError:
    #         os.remove(path)





    table2 = doc.add_table(rows=1, cols=20) 
    table2.style = 'resultados' 
    row = table2.rows[0].cells 
        

    fila = 1
    row[0].merge(row[0])
    row[0].paragraphs[0].add_run("No.").bold = True

    row[2].merge(row[1])
    row[1].paragraphs[0].add_run("FECHA").bold = True

    row[4].merge(row[3])
    row[3].paragraphs[0].add_run("RADICADO").bold = True

    row[6].merge(row[5])
    row[5].paragraphs[0].add_run("FOLIO N°").bold = True
                
    row[11].merge(row[7])
    row[7].paragraphs[0].add_run("ALERTA").bold = True


    row[14].merge(row[12])
    row[12].paragraphs[0].add_run("AMENAZA").bold = True


    row[19].merge(row[15])
    row[15].paragraphs[0].add_run("QUIEN RECIBE").bold = True
    

    for j in data:

        row = table2.add_row().cells 
        row[0].merge(row[0])
        row[0].text = str(fila)

        row[2].merge(row[1])
        row[1].text = str(j[2])

        row[4].merge(row[3])
        row[3].text = str(j[1])

        row[6].merge(row[5])
        row[5].text = str(j[29])
        
        row[11].merge(row[7])
        row[7].text = str(str(j[31])) 
        p = row[7].paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                
        row[14].merge(row[12])
        re_division = str(j[5]) + " - " + str(j[6]) + ". " + str(j[7])+ " - " + str(j[8]) + " (" + str(j[9])+ ") " 
        row[12].text = str(j[30]) 
        re_division_sub = str(j[33]) + " - " + str(j[34]) + ". " + str(j[35])+ " - " + str(j[37]) + " (" + str(j[36])+ ") " 
        quien_recibe =  str(re_division) +"\n" + "\n" + str(re_division_sub) 
                        
        row[19].merge(row[15])
        row[15].text = str(quien_recibe) 
        p = row[15].paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY



        # row[3].merge(row[0])
        # row[0].paragraphs[0].add_run("QUIEN RECIBE BRIGADA").bold = True
   
        # row[9].merge(row[4])
        # row[4].text = str(j[34])+ " - " +str(j[35]+". "+str(j[36]) + " (" + str(j[38])+")" + "  "+ str( j[37]))

        # row[13].merge(row[10])
        # row[10].paragraphs[0].add_run("QUIEN RECIBE BATALLÓN").bold = True
   
        # row[19].merge(row[14])
        # row[14].text = str(j[39])+ " - " +str(j[40]+". "+str(j[41]) + " (" + str(j[43])+")" + "  "+ str( j[42]))
        
        fila =  fila + 1

        parrafo = doc.add_paragraph()


        # row[6].merge(row[1])
        # row[1].text = str(data[j][3])
        # row[7].paragraphs[0].add_run("Tipo").bold = True
        # row[9].merge(row[8])
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()

    table3 = doc.add_table(rows=1, cols=10) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[9].merge(row[0])
    row[0].paragraphs[0].add_run(str(oficial_coe)).bold = True
    row = table3.add_row().cells 
    row[9].merge(row[0])
    row[0].paragraphs[0].add_run("OFICIAL CENTRO DE OPERACIONES DEL EJERCITO").bold = False

    doc.save(LINK+ "panilla alerta.docx")

    
    titulo = "panilla alerta"

    direcion= DIRECION+str(titulo)+'.docx'
    # print(direcion)
    return [direcion, titulo]

